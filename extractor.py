import pdfplumber
import json
import os
import re
import time
from dotenv import load_dotenv
from google import genai
from google.genai import types

load_dotenv()

try:
    client = genai.Client()
except Exception as e:
    pass

def _call_gemini_with_retry(prompt: str, max_retries: int = 3) -> str:
    for attempt in range(max_retries):
        try:
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt,
                config=types.GenerateContentConfig(response_mime_type="application/json"),
            )
            return response.text
        except Exception as e:
            error_msg = str(e)
            if attempt < max_retries - 1 and ("503" in error_msg or "429" in error_msg or "UNAVAILABLE" in error_msg):
                wait_time = 2 ** attempt  # Exponential backoff: 1s, 2s, 4s
                time.sleep(wait_time)
            else:
                raise e
    return None

def extract_text_from_pdf(pdf_path: str) -> str:
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    except Exception as e:
        return f"Error reading PDF: {str(e)}"
    return text


def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from DOCX files."""
    text = ""
    try:
        from docx import Document
        doc = Document(docx_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except ImportError:
        return "Error: python-docx library not installed. Please install it to support DOCX files."
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"
    return text


def extract_resume_text(file_path: str) -> str:
   
    if not file_path:
        return "Error: No file path provided"
    
    file_ext = file_path.rsplit('.', 1)[-1].lower() if '.' in file_path else ''
    
    if file_ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_ext in ('docx', 'doc'):
        return extract_text_from_docx(file_path)
    else:
        return f"Error: Unsupported file format: .{file_ext}"

def _extract_name_from_text(text: str) -> str:
    """Extract name from resume text with improved pattern matching."""
    if not text:
        return "Not provided"
    
    import re
    
    lines = text.strip().split('\n')
    
    for line in lines[:15]:
        line = line.strip()
        
        if not line or len(line) < 2:
            continue
        
        if len(line) > 60 or line.count('\t') > 2:
            continue
        
        if '@' in line or '.com' in line.lower() or 'http' in line.lower():
            continue
        
        headers = ['phone', 'email', 'linkedin', 'github', 'website', 'address', 'objective', 'summary', 
                  'experience', 'education', 'skills', 'certifications', 'projects', 'languages', 'references']
        if any(header in line.lower() for header in headers):
            continue
        
        # Check if it looks like a name (1-4 words, starts with capital)
        words = line.split()
        if 1 <= len(words) <= 4 and line[0].isupper() and not line.isupper():
            return line
    
    name_pattern = r'^([A-Z][a-z]+ [A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)$'
    for line in lines[:20]:
        line = line.strip()
        if re.match(name_pattern, line):
            return line
    
    email_match = re.search(r'([a-zA-Z0-9._%+-]+)@', text)
    if email_match:
        name_from_email = email_match.group(1).replace('.', ' ').replace('_', ' ').title()
        if len(name_from_email) > 2:
            return name_from_email
    
    return "Not provided"


def _extract_email_from_text(text: str) -> str:
    """Extract email from resume text with improved patterns."""
    if not text:
        return "Not provided"
    
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    if email_match:
        return email_match.group()
    
    alt_patterns = [
        r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
        r'(?:Email|email|EMAIL)[\s:]*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+)',
        r'(?:mailto:)?([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+)',
    ]
    
    for pattern in alt_patterns:
        match = re.search(pattern, text)
        if match:
            if '@' in match.group():
                return match.group().replace('mailto:', '')
    
    return "Not provided"


def _extract_years_experience_from_text(text: str) -> int:
    """Extract years of experience from resume text."""
    if not text:
        return 0
    
    experience_matches = re.findall(r'(\d+)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:experience|exp)', text.lower())
    if experience_matches:
        
        return max(int(x) for x in experience_matches)
    
    exp_match = re.search(r'(\d+)\s+years?\s+(?:of\s+)?(?:professional\s+)?experience', text.lower())
    if exp_match:
        return int(exp_match.group(1))
    
    return 0


def parse_resume_with_llm(text: str) -> dict:
    prompt = f"""
    You are an expert HR recruiter. Extract key information from this resume.
    Output ONLY a valid JSON object with these EXACT keys and ensure ALL fields are populated with non-null values:
    - "Name": Full name of candidate (string, MUST NOT be null or empty)
    - "Email": Email address (string, MUST NOT be null or empty) 
    - "Phone": Phone number if available (string, can be null)
    - "Skills": List of specific technical and professional skills (list of strings, minimum 5 items)
    - "Total_Years_Experience": Total years of professional experience (integer, minimum 0)
    
    IMPORTANT:
    - Extract SPECIFIC skills like "Python", "AWS", "Project Management", NOT vague terms like "Technical Skills"
    - Name: ALWAYS provide a name. If not clearly stated, extract from email, header, or context. Use "Not Provided" ONLY as absolute last resort.
    - Email: ALWAYS look for email addresses anywhere in text. If found, include it. Never return null.
    - Skills: Extract minimum 5-10 distinct skills. If fewer than 5 found, infer reasonable ones based on role.
    - Years: Look for any mention of years of experience. Default to 0 if not mentioned.
    - If any field cannot be populated, use a sensible string value instead of null.
    
    Resume Text:
    {text}
    """
    try:
        response_text = _call_gemini_with_retry(prompt)
        if response_text:
            parsed = json.loads(response_text)
            name = parsed.get('Name') or _extract_name_from_text(text)
            email = parsed.get('Email') or _extract_email_from_text(text)
            
            # Ensure non-empty values
            name = name if name and name.strip() else _extract_name_from_text(text)
            email = email if email and email.strip() else _extract_email_from_text(text)
            
            result = {
                'Name': name,
                'Email': email,
                'Phone': parsed.get('Phone'),
                'Skills': parsed.get('Skills') or [],
                'Total_Years_Experience': parsed.get('Total_Years_Experience') or 0
            }
            
            if not isinstance(result['Skills'], list):
                result['Skills'] = []
            
            if len(result['Skills']) == 0:
                skill_match = re.findall(r'(Python|Java|JavaScript|React|Node|AWS|Azure|Docker|Kubernetes|SQL|NoSQL|Git|Linux|Windows|Leadership|Communication|Problem[\s-]?Solving|Project Management|Agile|Scrum)', text, re.IGNORECASE)
                result['Skills'] = list(dict.fromkeys(skill_match))[:10]
            
            return result
        else:
            return {
                'Name': _extract_name_from_text(text),
                'Email': _extract_email_from_text(text),
                'Phone': None,
                'Skills': [],
                'Total_Years_Experience': _extract_years_experience_from_text(text)
            }
    except json.JSONDecodeError:
        return {
            'Name': _extract_name_from_text(text),
            'Email': _extract_email_from_text(text),
            'Phone': None,
            'Skills': [],
            'Total_Years_Experience': _extract_years_experience_from_text(text)
        }
    except Exception as e:
        return {
            'Name': _extract_name_from_text(text),
            'Email': _extract_email_from_text(text),
            'Phone': None,
            'Skills': [],
            'Total_Years_Experience': _extract_years_experience_from_text(text)
        }

def generate_feedback(resume_text: str, job_description: str) -> dict:
    prompt = f"""
    Act as an expert technical recruiter. Compare the candidate's resume to the job description.
    Identify the biggest gaps or missing keywords. 
    Provide specific, actionable feedback on how to improve the resume.
    Output ONLY a valid JSON object with a single key "improvements" containing a list of 5 specific improvement suggestions as strings.
    
    Format each suggestion starting with what to improve, then how to improve it.
    Example: "Add AWS certification to skills section - mention specific services used"
    
    Job Description:
    {job_description}
    
    Resume Text:
    {resume_text}
    """
    try:
        response_text = _call_gemini_with_retry(prompt)
        if response_text:
            result = json.loads(response_text)
            if "improvements" in result and isinstance(result["improvements"], list) and len(result["improvements"]) > 0:
                return result
            else:
                return {"improvements": ["Add more specific metric-driven achievements", "Highlight relevant certifications and training", "Optimize keywords for ATS systems", "Quantify your impact with numbers", "Tailor experience descriptions to job requirements"]}
        else:
            return {"improvements": ["API temporarily unavailable. Please try again in a moment."]}
    except json.JSONDecodeError:
        return {"improvements": ["Add more specific metric-driven achievements", "Highlight relevant certifications and training", "Optimize keywords for ATS systems", "Quantify your impact with numbers", "Tailor experience descriptions to job requirements"]}
    except Exception as e:
        return {"improvements": ["Add more achievement-oriented descriptions", "Include relevant technical skills", "Highlight measurable results and impact", "Use action verbs and keywords from job posting", "Format consistently for ATS compatibility"]}

def generate_template(job_description: str) -> dict:
    prompt = f"""
    Act as an expert resume writer. Based on the job description provided, generate a resume template with suggested sections and content.
    Output ONLY a valid JSON object with these EXACT keys: 
    - "summary" (2-3 sentence professional summary)
    - "key_sections" (list of 5-6 section names like "Professional Summary", "Experience", "Education", "Skills")
    - "suggested_bullets" (list of 5 achievement bullet points tailored to this role)
    
    Ensure ALL fields are non-empty lists or strings.
    
    Job Description:
    {job_description}
    """
    try:
        response_text = _call_gemini_with_retry(prompt)
        if response_text:
            result = json.loads(response_text)
            if "summary" in result and "key_sections" in result and "suggested_bullets" in result:
                if result["key_sections"] and result["suggested_bullets"]:
                    return result
        
        return _generate_template_fallback(job_description)
    except json.JSONDecodeError:
        return _generate_template_fallback(job_description)
    except Exception as e:
        return _generate_template_fallback(job_description)


def _generate_template_fallback(job_description: str) -> dict:
    jd_keywords = job_description.lower().split()
    

    is_senior = any(word in job_description.lower() for word in ["senior", "lead", "manager", "principal", "architect"])
    is_junior = any(word in job_description.lower() for word in ["junior", "entry", "graduate", "intern"])
    
    years_match = [w for w in jd_keywords if '+' in w and 'year' in job_description.lower()]
    
    if is_senior:
        summary = f"Results-driven professional with proven expertise in leading technical initiatives and delivering enterprise software solutions. Demonstrated success in building high-performing teams and driving organizational growth through innovation and strategic planning."
    elif is_junior:
        summary = f"Motivated professional with solid foundation in software development and problem-solving. Passionate about learning new technologies and contributing to collaborative team environments while building scalable solutions."
    else:
        summary = f"Experienced professional with expertise in software development, team collaboration, and delivering high-quality solutions. Committed to continuous learning and driving technical excellence across cross-functional initiatives."
    
    key_sections = [
        "Professional Summary",
        "Technical Skills",
        "Professional Experience",
        "Education",
        "Certifications & Achievements"
    ]
    
    if any(word in job_description.lower() for word in ["data", "analytics", "database"]):
        suggested_bullets = [
            "Developed and optimized data pipelines processing 100K+ records daily, improving query performance by 40%",
            "Designed scalable database architecture supporting 5M+ concurrent users",
            "Implemented analytics dashboards reducing reporting time by 60%",
            "Led data migration project from legacy systems saving $500K annually",
            "Mentored 3+ junior developers in database optimization best practices"
        ]
    elif any(word in job_description.lower() for word in ["cloud", "aws", "azure", "devops"]):
        suggested_bullets = [
            "Architected and deployed cloud infrastructure on AWS reducing operational costs by 35%",
            "Implemented CI/CD pipelines decreasing deployment time from 2 hours to 15 minutes",
            "Scaled applications to handle 10x traffic increase with zero downtime",
            "Managed containerization strategy using Docker and Kubernetes for 20+ microservices",
            "Reduced infrastructure incidents by 70% through automated monitoring and alerting"
        ]
    elif any(word in job_description.lower() for word in ["frontend", "react", "vue", "angular", "ui"]):
        suggested_bullets = [
            "Built responsive user interfaces with React serving 500K+ daily active users",
            "Improved page load time by 50% through code splitting and lazy loading optimization",
            "Designed and implemented component library increasing dev velocity by 30%",
            "Led frontend accessibility initiative ensuring WCAG 2.1 AA compliance",
            "Mentored team of 4 frontend developers on best practices and code quality standards"
        ]
    else:
        suggested_bullets = [
            "Led development of mission-critical features delivering 30% improvement in user engagement",
            "Architected scalable solutions supporting 10x business growth with 99.9% uptime",
            "Implemented automated testing framework increasing code coverage from 40% to 85%",
            "Collaborated cross-functionally with product, design, and leadership teams to deliver quarterly roadmap",
            "Mentored junior team members and conducted code reviews ensuring quality standards"
        ]
    
    return {
        "summary": summary,
        "key_sections": key_sections,
        "suggested_bullets": suggested_bullets
    }

def categorize_missing_skills(missing_skills: list, job_description: str, present_skills: list) -> dict:
    """Categorize missing skills by priority and provide development recommendations."""
    
    prompt = f"""
    You are an expert career development coach. Categorize these missing skills by priority and provide learning recommendations.
    
    MISSING SKILLS TO CATEGORIZE: {', '.join(missing_skills[:8])}
    SKILLS CANDIDATE ALREADY HAS: {', '.join(present_skills[:10])}
    
    For each missing skill, determine:
    1. Priority level: "Critical" (essential), "High" (important), "Medium" (valuable), "Nice-to-have"
    2. Learning time: How long to acquire (e.g., "2-4 weeks", "3-6 months")
    3. Learning path: Best way to learn (e.g., "Online course", "Practice project", "Certification")
    4. Impact: How much it affects hiring chances (1-10 scale)
    
    Output ONLY a valid JSON object with this structure:
    {{
        "critical_skills": [
            {{"skill": "Skill Name", "timeframe": "2-4 weeks", "learning_path": "Online course", "impact": 9}}
        ],
        "high_priority_skills": [
            {{"skill": "Skill Name", "timeframe": "4-8 weeks", "learning_path": "Project-based", "impact": 8}}
        ],
        "medium_priority_skills": [
            {{"skill": "Skill Name", "timeframe": "2-3 months", "learning_path": "Certification course", "impact": 6}}
        ],
        "nice_to_have_skills": [
            {{"skill": "Skill Name", "timeframe": "Ongoing", "learning_path": "Self-study", "impact": 4}}
        ]
    }}
    
    Job Description for context:
    {job_description}
    """
    
    try:
        response_text = _call_gemini_with_retry(prompt)
        if response_text:
            categorized = json.loads(response_text)
            return {
                "categorized_skills": categorized,
                "total_missing": len(missing_skills),
                "immediate_action_skills": categorized.get("critical_skills", [])[:3]
            }
    except Exception:
        pass
    
    # Fallback categorization using keyword matching
    skill_keywords = {
        "critical": ["Python", "Java", "AWS", "Docker", "Kubernetes", "SQL", "Leadership", "API", "Git"],
        "high": ["JavaScript", "React", "Angular", "Node", "DevOps", "Testing", "CI/CD", "Database"],
        "medium": ["CSS", "HTML", "Agile", "Communication", "Project Management", "Version Control"],
        "nice": ["Presentation", "Documentation", "Writing", "Public Speaking"]
    }
    
    categorized = {
        "critical_skills": [],
        "high_priority_skills": [],
        "medium_priority_skills": [],
        "nice_to_have_skills": []
    }
    
    for skill in missing_skills:
        skill_lower = skill.lower()
        assigned = False
        
        for level, keywords in skill_keywords.items():
            if any(kw.lower() in skill_lower for kw in keywords):
                if level == "critical":
                    categorized["critical_skills"].append({"skill": skill, "timeframe": "2-4 weeks", "learning_path": "Online course + projects", "impact": 9})
                elif level == "high":
                    categorized["high_priority_skills"].append({"skill": skill, "timeframe": "4-8 weeks", "learning_path": "Projects + documentation", "impact": 8})
                elif level == "medium":
                    categorized["medium_priority_skills"].append({"skill": skill, "timeframe": "2-3 months", "learning_path": "Courses + practice", "impact": 6})
                else:
                    categorized["nice_to_have_skills"].append({"skill": skill, "timeframe": "Ongoing", "learning_path": "Self-study", "impact": 4})
                assigned = True
                break
        
        if not assigned:
            categorized["high_priority_skills"].append({"skill": skill, "timeframe": "4-8 weeks", "learning_path": "Research + practice", "impact": 7})
    
    return {
        "categorized_skills": categorized,
        "total_missing": len(missing_skills),
        "immediate_action_skills": categorized["critical_skills"][:3]
    }


def analyze_skill_gaps(resume_text: str, job_description: str) -> dict:
    """Extract and categorize skill gaps using Gemini AI with detailed analysis."""
    
    extraction_prompt = f"""
    You are an expert technical recruiter with 20+ years of experience. Your task is to extract and compare skills.
    
    TASK 1 - EXTRACT FROM RESUME:
    List EVERY technical skill, programming language, tool, framework, platform, and professional competency mentioned in the resume text below.
    Be exhaustive - include languages, databases, frameworks, methodologies, soft skills, certifications, etc.
    
    TASK 2 - EXTRACT FROM JOB DESCRIPTION:
    List EVERY skill and requirement mentioned in the job description. Include required skills, preferred skills, nice-to-have skills.
    
    TASK 3 - IDENTIFY GAPS:
    Find skills that appear in the job description but NOT in the resume.
    
    TASK 4 - CATEGORIZE BY IMPORTANCE:
    For each missing skill, determine if it's:
    - "Critical" (mentioned multiple times, core to the role)
    - "Important" (core responsibility or requirement)
    - "Valuable" (nice-to-have or supplementary)
    
    IMPORTANT GUIDELINES:
    - Be SPECIFIC: "Python" not "Programming"; "AWS Lambda" not "Cloud"; "React" not "Frontend"
    - Extract minimum 15-20 skills per list for comprehensive analysis
    - Include soft skills: "Leadership", "Communication", "Problem Solving", "Team Collaboration"
    - NEVER use vague terms like "Technical Skills", "Experience", "Knowledge"
    
    Return ONLY valid JSON (no markdown, no text before/after):
    {{
        "resume_skills": [list of 15-25 specific skills from resume],
        "job_description_skills": [list of 15-25 specific skills from job posting],
        "critical_missing_skills": [skills required but missing from resume],
        "important_missing_skills": [supplementary skills missing from resume],
        "valuable_missing_skills": [nice-to-have skills missing from resume]
    }}
    
    Resume Text:
    {resume_text}
    
    Job Description:
    {job_description}
    """
    
    try:
        response_text = _call_gemini_with_retry(extraction_prompt)
        if response_text:
            result = json.loads(response_text)
            
            # Validate response structure
            resume_skills = result.get("resume_skills", [])
            job_skills = result.get("job_description_skills", [])
            critical_missing = result.get("critical_missing_skills", [])
            important_missing = result.get("important_missing_skills", [])
            valuable_missing = result.get("valuable_missing_skills", [])
            
            # Combine all missing skills
            all_missing = critical_missing + important_missing + valuable_missing
            
            # Stage 2: Get learning recommendations for missing skills
            if all_missing:
                recommendations = _get_skill_development_plan(all_missing, job_description, resume_skills)
            else:
                recommendations = {"categorized_skills": {}, "learning_paths": {}}
            
            return {
                "present_skills": resume_skills[:20] if resume_skills else ["Communication", "Problem Solving"],
                "required_skills": job_skills[:20] if job_skills else ["Technical skills"],
                "missing_skills": critical_missing + important_missing + valuable_missing,
                "missing_skills_by_priority": {
                    "critical": critical_missing[:5],
                    "important": important_missing[:5],
                    "valuable": valuable_missing[:5]
                },
                "skill_development_plan": recommendations
            }
    except json.JSONDecodeError as e:
        print(f"JSON parsing error in skill extraction: {e}")
    except Exception as e:
        print(f"Error in skill gap analysis: {e}")
    
    # Fallback to simpler extraction
    return _fallback_skill_gap_analysis(resume_text, job_description)


def _get_skill_development_plan(missing_skills: list, job_description: str, present_skills: list) -> dict:
    """Get learning recommendations for missing skills using Gemini."""
    
    skills_list = ", ".join(missing_skills[:15])
    present_list = ", ".join(present_skills[:10])
    
    plan_prompt = f"""
    You are a career development expert. Create a learning and development plan for acquiring these missing skills.
    
    MISSING SKILLS TO DEVELOP: {skills_list}
    SKILLS ALREADY POSSESSED: {present_list}
    JOB ROLE CONTEXT: {job_description[:500]}
    
    For EACH missing skill, provide:
    1. Priority: 1-10 (10 = most critical for getting hired)
    2. Learning timeframe: Realistic time to gain proficiency
    3. Best learning methods: Specific resources, courses, projects, certifications
    4. How to demonstrate: How to show this skill to employers (portfolio, project, certification, GitHub)
    
    Return ONLY valid JSON:
    {{
        "skill_development": [
            {{
                "skill": "Skill Name",
                "priority": 9,
                "timeframe": "4-6 weeks",
                "learning_methods": ["Udemy course XYZ", "Build 2-3 projects", "Practice with datasets"],
                "demonstration": "GitHub portfolio + completed projects",
                "resources": ["URL 1", "URL 2"]
            }}
        ],
        "quick_wins": ["Skills that can be picked up in 1-2 weeks"],
        "long_term_goals": ["Skills requiring 3+ months of dedicated learning"],
        "estimated_total_learning_time": "3-6 months"
    }}
    """
    
    try:
        response_text = _call_gemini_with_retry(plan_prompt)
        if response_text:
            plan = json.loads(response_text)
            return plan
    except Exception as e:
        print(f"Error getting skill development plan: {e}")
    
    return {"categorized_skills": {}, "learning_paths": {}}


def _fallback_skill_gap_analysis(resume_text: str, job_description: str) -> dict:
    """Fallback skill gap analysis when Gemini fails."""
    
    skill_patterns = {
        "programming": r"\b(Python|Java|JavaScript|TypeScript|C\+\+|C#|Go|Rust|Ruby|PHP|Swift|Kotlin|R|MATLAB|Scala)\b",
        "web": r"\b(React|Vue|Angular|Node\.js|Express|Django|Flask|Spring|FastAPI|Next\.js|Svelte|HTML|CSS|SASS|Bootstrap)\b",
        "cloud": r"\b(AWS|Azure|Google Cloud|GCP|Lambda|EC2|S3|RDS|Firebase|Heroku|DigitalOcean)\b",
        "devops": r"\b(Docker|Kubernetes|Jenkins|GitHub Actions|GitLab CI|Terraform|Ansible|CloudFormation|CI/CD)\b",
        "database": r"\b(SQL|PostgreSQL|MySQL|MongoDB|Redis|Elasticsearch|DynamoDB|Oracle|Firebase|Cassandra)\b",
        "soft": r"\b(Leadership|Communication|Problem[\s-]?Solving|Team[\s-]?Work|Project[\s-]?Management|Agile|Scrum|Presentation)\b"
    }
    
    resume_skills = []
    job_skills = []
    
    for category, pattern in skill_patterns.items():
        resume_matches = re.findall(pattern, resume_text, re.IGNORECASE)
        job_matches = re.findall(pattern, job_description, re.IGNORECASE)
        resume_skills.extend(resume_matches)
        job_skills.extend(job_matches)
    
    resume_skills = list(dict.fromkeys(resume_skills))  # Remove duplicates
    job_skills = list(dict.fromkeys(job_skills))
    
    missing = [skill for skill in job_skills if skill not in resume_skills]
    
    return {
        "present_skills": resume_skills[:20] if resume_skills else ["Communication", "Problem Solving"],
        "required_skills": job_skills[:20] if job_skills else ["Technical skills"],
        "missing_skills": missing[:15] if missing else ["Specialized skills"],
        "missing_skills_by_priority": {
            "critical": missing[:5] if missing else [],
            "important": missing[5:10] if len(missing) > 5 else [],
            "valuable": missing[10:] if len(missing) > 10 else []
        },
        "skill_development_plan": {}
    }